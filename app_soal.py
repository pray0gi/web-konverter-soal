import streamlit as st
import docx
import re
import urllib.parse
import io
import base64
import requests
from docx.oxml.ns import qn

# --- FUNGSI PARSER (GAMBAR LANGSUNG KE RAM/BASE64) ---
def ekstrak_gambar_base64(paragraph, doc):
    """Mengekstrak gambar dari paragraf Word dan mengubahnya menjadi string Base64"""
    list_b64 = []
    blips = paragraph._element.xpath('.//a:blip')
    for blip in blips:
        rId = blip.get(qn('r:embed'))
        if rId and rId in doc.part.related_parts:
            image_part = doc.part.related_parts[rId]
            image_data = image_part.blob
            mime_type = image_part.content_type
            
            # Encode biner gambar ke teks Base64
            b64_str = base64.b64encode(image_data).decode('utf-8')
            list_b64.append({"mimeType": mime_type, "data": b64_str})
    return list_b64

def proses_konversi_api(file_upload):
    """Membaca isi Word, memotong nomor soal, dan menyusun JSON List"""
    doc = docx.Document(file_upload)
    data_soal = []
    
    soal_sementara = {"Pertanyaan": "", "Gambar_Pertanyaan": [], "A": "", "B": "", "C": "", "D": "", "E": "", "Kunci": ""}
    id_soal = 0
    
    for para in doc.paragraphs:
        teks = para.text.strip()
        miliki_gambar = len(para._element.xpath('.//a:blip')) > 0
        
        if not teks and not miliki_gambar:
            continue
            
        # 1. Deteksi Soal Baru
        if re.match(r'^\d+\.', teks):
            if soal_sementara["Pertanyaan"]:
                data_soal.append(soal_sementara.copy())
            id_soal += 1
            soal_sementara = {"Pertanyaan": "", "Gambar_Pertanyaan": [], "A": "", "B": "", "C": "", "D": "", "E": "", "Kunci": ""}
            
            # LOGIKA BARU: Hapus angka dan titik di depan pertanyaan (misal "1. ")
            teks_bersih = re.sub(r'^\d+\.\s*', '', teks)
            soal_sementara["Pertanyaan"] = teks_bersih.strip()
            
            if miliki_gambar:
                soal_sementara["Gambar_Pertanyaan"].extend(ekstrak_gambar_base64(para, doc))
                
        # 2. Deteksi Opsi Jawaban
        elif re.match(r'^[A-E]\.', teks, re.IGNORECASE):
            huruf = teks[0].upper()
            soal_sementara[huruf] = teks[2:].strip()
            
        # 3. Deteksi Kunci Jawaban
        elif teks.lower().startswith("kunci:"):
            soal_sementara["Kunci"] = teks.split(":")[1].strip().upper()
            
        # 4. Handle teks atau gambar sambungan di bawah pertanyaan
        elif id_soal > 0 and not re.match(r'^[A-E]\.', teks, re.IGNORECASE) and not teks.lower().startswith("kunci:"):
            if teks:
                soal_sementara["Pertanyaan"] += "\n" + teks
            if miliki_gambar:
                soal_sementara["Gambar_Pertanyaan"].extend(ekstrak_gambar_base64(para, doc))

    # Masukkan soal terakhir ke dalam list
    if soal_sementara["Pertanyaan"]:
        data_soal.append(soal_sementara)
        
    return data_soal

# --- ANTARMUKA STREAMLIT ---
st.set_page_config(page_title="Konverter Soal Ujian", layout="wide")

# --- MENU SIDEBAR: DOWNLOAD TEMPLATE ---
with st.sidebar:
    st.header("📄 Format Standar Soal")
    st.markdown("Agar aplikasi dapat membaca soal dengan benar, pastikan Anda menggunakan template Word yang telah disediakan.")
    
    def buat_template_memori():
        doc = docx.Document()
        doc.add_heading('Template Format Soal Ujian', 0)
        doc.add_paragraph('Petunjuk: Gunakan format di bawah ini dengan presisi. Jangan ubah pola penomoran (1., 2.) atau format huruf opsi (A., B.). Ketik manual, jangan gunakan fitur Auto-Numbering MS Word.')
        doc.add_paragraph('')
        
        doc.add_paragraph('1. Pada pemrograman PLC, alamat memori yang umum digunakan untuk input fisik dari perangkat luar adalah...')
        doc.add_paragraph('A. CIO Area')
        doc.add_paragraph('B. HR Area')
        doc.add_paragraph('C. DM Area')
        doc.add_paragraph('D. AR Area')
        doc.add_paragraph('E. Timer/Counter Area')
        doc.add_paragraph('Kunci: A')
        doc.add_paragraph('')
        
        doc.add_paragraph('2. Perhatikan gambar ladder diagram berikut ini.\n[SISIPKAN GAMBAR DI SINI]\nKomponen instruksi yang ditunjukkan oleh alamat 0.00 berfungsi sebagai...')
        doc.add_paragraph('A. Output / Coil')
        doc.add_paragraph('B. Normally Open (NO) Contact')
        doc.add_paragraph('C. Normally Closed (NC) Contact')
        doc.add_paragraph('D. Timer')
        doc.add_paragraph('E. Internal Relay')
        doc.add_paragraph('Kunci: B')
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    file_template = buat_template_memori()
    
    st.download_button(
        label="⬇️ Unduh Template Word (.docx)",
        data=file_template,
        file_name="Template_Soal_Ujian.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary",
        use_container_width=True
    )
    st.write("---")

# --- KONTEN UTAMA ---
st.title("📝 Web Konverter Soal Word ke Form")
st.markdown("Aplikasi untuk mengubah format soal ujian `.docx` menjadi Google Form secara otomatis, mendukung ekstraksi gambar dan acak soal.")

file_word = st.file_uploader("Unggah File Word (.docx)", type=["docx"])

if file_word is not None:
    st.write("---") 
    
    col1, col2 = st.columns(2)
    with col1:
        nama_form = st.text_input("Nama File Form:", "Soal Ujian Teknik Ketenagalistrikan")
        poin_soal = st.number_input("Poin per soal benar:", min_value=1, max_value=100, value=10)
    with col2:
        nomor_wa = st.text_input("Nomor WA Tujuan (Opsional):", placeholder="Contoh: 081234567890")
    
    st.write("") 
    tombol_konversi = st.button("🚀 BUAT GOOGLE FORM SEKARANG", type="primary", use_container_width=True)
    
    if tombol_konversi:
        with st.spinner("Sedang membaca dokumen, menyandikan gambar, dan mengirim ke Google..."):
            
            # Panggil fungsi parser Word -> JSON List
            data_json = proses_konversi_api(file_word) 
            
            # Bungkus menjadi payload API
            payload_api = {
                "poin": poin_soal,
                "nama_form": nama_form,
                "soal": data_json
            }
            
            # ⚠️ GANTI DENGAN URL WEB APP GAS ANDA
            URL_GAS = "https://script.google.com/macros/s/AKfycbxRUyFNd37oae91glDYewdP3-TFggDysiG1nSNOGnkXbGot0LHI9TmhSe8QKgLk8_Fgpw/exec"
            
            try:
                respon = requests.post(URL_GAS, json=payload_api)
                hasil_api = respon.json()
                
                if hasil_api.get("status") == "sukses":
                    st.success(f"🎉 Berhasil! Form '{nama_form}' sudah siap.")
                    
                    link_edit = hasil_api['edit_url']
                    link_siswa = hasil_api['publish_url']
                    
                    st.markdown(f"**🔗 [Klik di sini untuk Edit Form]({link_edit})**")
                    st.markdown(f"**👀 [Klik di sini untuk Lihat Tampilan Form]({link_siswa})**")
                    
                    # LOGIKA WHATSAPP
                    if nomor_wa:
                        pesan_wa = f"Halo! Google Form untuk *{nama_form}* sudah selesai dibuat.\n\n*Link Edit (Guru):*\n{link_edit}\n\n*Link Kuis (Siswa):*\n{link_siswa}"
                        
                        # Format nomor 08 menjadi 628
                        if nomor_wa.startswith("08"):
                            nomor_wa = "628" + nomor_wa[2:] 
                        
                        pesan_terencode = urllib.parse.quote(pesan_wa)
                        link_wa_web = f"https://wa.me/{nomor_wa}?text={pesan_terencode}"
                        
                        st.info("Form berhasil dibuat. Klik tombol di bawah ini untuk meneruskan link via WhatsApp.")
                        st.markdown(f"**📱 [KIRIM LINK VIA WHATSAPP]({link_wa_web})**")
                        
                else:
                    st.error(f"Terjadi kesalahan di server Google: {hasil_api.get('pesan')}")
                    
            except Exception as e:
                st.error(f"Gagal menghubungi server: {e}. Pastikan URL GAS sudah dimasukkan dengan benar.")
